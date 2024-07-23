from docx import Document
import locale
from docx.shared import Pt
import re
from decimal import Decimal, InvalidOperation

locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')

def format_number(price):
    european_pattern = re.compile(r'^\d{1,3}(\.\d{3})*,\d{2}$')
    american_pattern = re.compile(r'^\d{1,3}(,\d{3})*\.\d{2}$')

    if european_pattern.match(price):
        price = price.replace('.', '').replace(',', '.')
    elif american_pattern.match(price):
        price = price.replace(',', '')
    else:
        return "Invalid number"

    try:
        number_decimal = Decimal(price)
        formatted_number = locale.currency(number_decimal, grouping=True, symbol=False)
        return formatted_number
    except InvalidOperation:
        return "Invalid number"

def apply_font_size_content(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(8)

def apply_font_size_title(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = True

def write_tables(document: Document, df) -> Document:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'TableGrid'

    header_cells = table.rows[0].cells
    for i in range(len(df.columns)):
        try:
            text = df.columns[i]
            if "Unnamed" in str(text):
                text = ""
            header_cells[i].text = text
            apply_font_size_title(header_cells[i])
        except Exception as ex:
            print(ex)

    number_lines = len(df)
    number_columns = len(df.columns)

    amount_total = Decimal(0)
    
    for line in range(number_lines):
        row_cells = table.add_row().cells
        for column in range(number_columns):
            try:
                text = str(df[df.columns[column]][line])
                if str(text) == 'nan':
                    text = ''
                
                if df.columns[column] == 'Unit price' or df.columns[column] == 'Amount':
                    text = format_number(text)
                    print("Amount: ", text)
                    if df.columns[column] == 'Amount':
                        number = text.replace(',', '').replace('$', '').strip()
                        amount_total += Decimal(number)
                
                row_cells[column].text = text
                apply_font_size_content(row_cells[column])
            except Exception as ex:
                print(ex)

    formatted_total = format_number(str(amount_total))
    text_result = ['Total Amount:', formatted_total]
    print(text_result)

    row_cells = table.add_row().cells
    for i in range(2):
        try:
            text = text_result[i]
            if str(text) == 'nan':
                text = ''
            index = (number_columns - 2) + i
            row_cells[index].text = text
            apply_font_size_content(row_cells[index])
        except Exception as ex:
            print("error result line", ex)

    return document