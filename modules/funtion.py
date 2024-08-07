import customtkinter as ctk
import datetime
from tkinter import filedialog
from tkinter.scrolledtext import ScrolledText
from customtkinter import CTkImage
from docx import Document
import fitz
import pyautogui
from PIL import Image, ImageDraw, ImageFont

data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")

def exibir_pop_up(mensagem):
    popup = ctk.CTk()
    popup.title("Erro")
    label = ctk.CTkLabel(popup, text=mensagem, width=300)
    label.pack(padx=20, pady=20)
    button = ctk.CTkButton(popup, text="OK", command=popup.destroy)
    button.pack(pady=10)
    popup.mainloop()

def carregar_pdf(pdf_entry, log_text, check_files):
    global pdf_path
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if pdf_path:
        pdf_entry.delete(0, ctk.END)
        pdf_entry.insert(0, pdf_path)
        data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
        log_text.configure(state="normal")
        log_text.insert(ctk.END, f"{data_hora_atual} - PDF importado com sucesso...\n", "log")
        log_text.configure(state="disabled")
        check_files()

def extrair_texto_pdf(pdf_path, texto_procurado):
    documento = fitz.open(pdf_path)
    for pagina in documento:
        texto = pagina.get_text()
        if texto_procurado in texto:
            posicao_inicio = texto.find(texto_procurado)
            if posicao_inicio != -1:
                posicao_fim = texto.find(" ", posicao_inicio + len(texto_procurado) + 1)
                texto_extraido = texto[posicao_inicio + len(texto_procurado) + 1:posicao_fim].strip()
                return texto_extraido.split()[0]
    return None

def processar(log_text):
    data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
    log_text.configure(state="normal")
    log_text.insert(ctk.END, f"{data_hora_atual} - Processamento iniciado...\n", "log")
    log_text.configure(state="disabled")
    log_text.configure(state="normal")
    log_text.insert(ctk.END, f"{data_hora_atual} - Arquivos processados com sucesso!\n", "log")
    log_text.configure(state="disabled")
 
def criar_documento_word(texto, requisicao_id, file_path):
    data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
    doc = Document()
    doc.add_heading("Relatório de Processamento", 0)
    doc.add_paragraph('Este é um relatório de exemplo.')
    doc.add_paragraph(f'Import PO No: {texto}')
    doc.add_paragraph(f'ID da Requisição: {requisicao_id}')
    doc.add_paragraph(f'Data e Hora de criação: {data_hora_atual}')
    doc.save(file_path)

def export_report(log_text, requisicao_id, pdf_path):
    data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
    log_text.configure(state="normal")
    log_text.insert(ctk.END, f"{data_hora_atual} - Exportação do relatório iniciada...\n", "log")
    log_text.configure(state="disabled")

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])

    if file_path:
        texto = extrair_texto_pdf(pdf_path, "Import PO No")
        if texto:
            criar_documento_word(texto, requisicao_id, file_path)
            
            log_text.configure(state="normal")
            log_text.insert(ctk.END, f"{data_hora_atual} - Relatório exportado com sucesso!\n", "log")
            log_text.configure(state="disabled")
        else:
            log_text.configure(state="normal")
            log_text.insert(ctk.END, f"{data_hora_atual} - Erro ao Exportar Relatório!\n", "log")
            log_text.configure(state="disabled")
            exibir_pop_up("Erro: O texto 'Import PO No' não foi encontrado no PDF.")
        
        save_log_to_image(log_text)

def process_export(log_text, requisicao_id_entry, pdf_entry):
    id_value = requisicao_id_entry.get()
    pdf_value = pdf_entry.get()

    id_length_valid = len(id_value) == 5 and not is_placeholder(requisicao_id_entry, "Digite o ID")
    pdf_filled = bool(pdf_value) and not is_placeholder(pdf_entry, "Selecione o PDF")

    if id_length_valid and pdf_filled:
        processar(log_text)
        export_report(log_text, id_value, pdf_value)
    else:
        data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
        log_text.configure(state="normal")
        log_text.insert(ctk.END, f"{data_hora_atual} - Entradas inválidas para processamento e/ou exportação.\n", "log")
        log_text.configure(state="disabled")

def only_numbers(char, current_text):
    return char.isdigit() and len(current_text) < 6

def add_placeholder(entry, placeholder, color="grey"):
    def on_focus_in(event):
        if is_placeholder(entry, placeholder):
            entry.delete(0, "end")
            entry.configure(text_color="black")

    def on_focus_out(event):
        if not entry.get():
            entry.insert(0, placeholder)
            entry.configure(text_color=color)

    if not entry.get():
        entry.insert(0, placeholder)
        entry.configure(text_color=color)

    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)

def is_placeholder(entry, placeholder):
    return entry.get() == placeholder

def adicionar_id(requisicao_id_entry, log):
    data_hora_atual = datetime.datetime.now().strftime("%d-%m-%y %H:%M:%S")
    id_value = requisicao_id_entry.get()
    log.configure(state="normal")
    log.insert("end", f"{data_hora_atual} - ID Adicionado: {id_value}\n", "log")
    log.configure(state="disabled")

def validate_entries(pdf_entry, requisicao_id_entry, adicionar_id_btn, process_export_btn):
    id_value = requisicao_id_entry.get()
    pdf_value = pdf_entry.get()

    id_length_valid = len(id_value) == 5 and not is_placeholder(requisicao_id_entry, "Digite o ID")
    pdf_filled = bool(pdf_value) and not is_placeholder(pdf_entry, "Selecione o PDF")

    if id_length_valid:
        adicionar_id_btn.configure(state=ctk.NORMAL)
    else:
        adicionar_id_btn.configure(state=ctk.DISABLED)

    if id_length_valid and pdf_filled:
        process_export_btn.configure(state=ctk.NORMAL)
    else:
        process_export_btn.configure(state=ctk.DISABLED)

def capture_log_screenshot(widget):
    widget.update_idletasks()
    x = widget.winfo_rootx()
    y = widget.winfo_rooty()
    width = widget.winfo_width()
    height = widget.winfo_height()

    screenshot = pyautogui.screenshot(region=(x, y, width, height))

    screenshot.save("log_screenshot.png")

def save_log_to_file(log_text):
    log_content = log_text.get("1.0", ctk.END)
    with open("log_content.txt", "w") as file:
        file.write(log_content)

def save_log_to_image(log_text):
    log_content = log_text.get("1.0", ctk.END)
    
    font = ImageFont.truetype("arial.ttf", 14)
    max_width = 800
    line_height = font.getbbox("A")[3] - font.getbbox("A")[1]
    
    lines = log_content.splitlines()
    spacing = 20
    img_height = (line_height + spacing) * len(lines) + 20

    img = Image.new("RGB", (max_width, img_height), "white")
    draw = ImageDraw.Draw(img)

    y = 10
    for line in lines:
        draw.text((10, y), line, fill="black", font=font)
        y += line_height + spacing

    img.save("log_content.png")
