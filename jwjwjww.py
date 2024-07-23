from docx import Document
import locale
from docx.shared import Pt
import re
from decimal import Decimal, InvalidOperation

locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')

def format_number(price):
    # Padrão de formato europeu (1.234,56)
    european_pattern = re.compile(r'^\d{1,3}(\.\d{3})*,\d{2}$')
    # Padrão de formato americano (1,234.56)
    american_pattern = re.compile(r'^\d{1,3}(,\d{3})*\.\d{2}$')

    if european_pattern.match(price):
        # Substituir pontos por nada e vírgulas por pontos
        price = price.replace('.', '').replace(',', '.')
    elif american_pattern.match(price):
        # Substituir vírgulas por nada
        price = price.replace(',', '')
    else:
        return "Invalid number"

    try:
        # Convertendo a string para Decimal para manter a precisão
        number_decimal = Decimal(price)

        # Formatando o número como moeda americana
        formatted_number = locale.currency(number_decimal, grouping=True, symbol=True)
        return formatted_number
    except InvalidOperation:
        return "Invalid number"

def apply_font_size_content(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(8)

def apply_font_size_title(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = True

def write_tables(document: Document, df) -> Document:
    """
        This function writes tables in the document.

        Params:
            document: Document -> Expected document import PO
            df: DataFrame -> DataFrame with data to write in the document

        Returns:
            Document with tables
    """

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'TableGrid'

    header_cells = table.rows[0].cells
    # header
    for i in range(len(df.columns)):
        try:
            text = df.columns[i]
            if "Unnamed" in str(text):
                text = ""
            header_cells[i].text = text
            apply_font_size_title(header_cells[i])
        except Exception as ex:
            print(ex)

    number_lines = len(df)
    number_columns = len(df.columns)

    # content
    for line in range(number_lines):
        row_cells = table.add_row().cells
        for column in range(number_columns):
            try:
                text = str(df[df.columns[column]][line])

                if str(text) == 'nan':
                    text = ''

                if df.columns[column] == 'Unit price' or df.columns[column] == 'Amount':
                    text = format_number(text)
                    print("Unit price: ", text)

                row_cells[column].text = text
                apply_font_size_content(row_cells[column])
            except Exception as ex:
                print(ex)

    # footer
    amount_list = []
    for i in range(len(df['Amount'].values)):
        number = df['Amount'].values[i]
        if type(number) == type(None) or number == 'nan':
            continue
        number = number.replace('.', '').replace(',', '.')
        try:
            amount_list.append(Decimal(number))
        except InvalidOperation:
            print(f"Invalid number in Amount: {number}")
            continue

    amount_total = sum(amount_list)
    print("Total", amount_total)

    amount_total_as_string = str(amount_total)
    formatted_total = format_number(amount_total_as_string)
    text_result = ['Total Amount:', formatted_total]
    print(text_result)

    row_cells = table.add_row().cells
    for i in range(2):
        try:
            text = text_result[i]
            if str(text) == 'nan':
                text = ''
            index = (number_columns - 2) + i
            row_cells[index].text = text
            apply_font_size_content(row_cells[index])
        except Exception as ex:
            print("error result line", ex)

    return document